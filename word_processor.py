import os
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import docx
from docx import Document as DocxDocument
import zipfile

def process_word(file_path):
    """
    Loads a Word document (.docx), extracts text, and splits it into chunks.
    Returns empty list if Word processing fails.
    """
    try:
        # Check if file is a valid .docx file
        if not is_valid_docx(file_path):
            print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            print("ERROR: The file is not a valid .docx document.")
            print("Please ensure you're uploading a Microsoft Word .docx file.")
            print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            return []

        # Load the Word document
        doc = DocxDocument(file_path)
        
        # Extract text from all paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text.strip())
        
        # Extract text from tables if any
        table_text = extract_tables_text(doc)
        if table_text:
            full_text.extend(table_text)
        
        # Join all text
        document_text = '\n\n'.join(full_text)
        
        # Check if any text was extracted
        if not document_text.strip():
            print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            print("ERROR: No text was extracted from the Word document.")
            print("The document might be empty or contain only images/objects.")
            print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            return []
        
        print("--- WORD PROCESSING DEBUG ---")
        print(f"Successfully extracted text from Word document.")
        print(f"Total paragraphs processed: {len([p for p in doc.paragraphs if p.text.strip()])}")
        print(f"Total tables processed: {len(doc.tables)}")
        print(f"Total text length: {len(document_text)} characters")
        preview_text = document_text[:300].replace('\n', ' ').strip()
        print(f"First 300 characters preview: \n'{preview_text}...'")
        print("-----------------------------")
        
        # Create a document object
        documents = [Document(
            page_content=document_text,
            metadata={"source": file_path, "type": "word_document"}
        )]
        
        # Use text splitter for better processing
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,  # Same as PDF for consistency
            chunk_overlap=300,  # Good overlap for context preservation
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]  # Hierarchical separation
        )
        
        texts = text_splitter.split_documents(documents)
        
        print(f"Created {len(texts)} text chunks for processing.")
        
        return texts
        
    except Exception as e:
        print(f"Error processing Word document: {str(e)}")
        return []

def is_valid_docx(file_path):
    """
    Check if the file is a valid .docx file by examining its structure.
    """
    try:
        # Check if it's a valid zip file (docx files are zip archives)
        with zipfile.ZipFile(file_path, 'r') as zip_file:
            # Check for required docx components
            required_files = ['[Content_Types].xml', 'word/document.xml']
            file_list = zip_file.namelist()
            
            for required_file in required_files:
                if required_file not in file_list:
                    return False
            
            return True
    except:
        return False

def extract_tables_text(doc):
    """
    Extract text from all tables in the Word document.
    """
    table_texts = []
    
    try:
        for table_idx, table in enumerate(doc.tables):
            table_content = []
            table_content.append(f"TABLE {table_idx + 1}:")
            
            # Extract headers if present
            if table.rows:
                header_row = table.rows[0]
                headers = []
                for cell in header_row.cells:
                    headers.append(cell.text.strip())
                
                if any(headers):  # If headers contain text
                    table_content.append("Headers: " + " | ".join(headers))
            
            # Extract all rows
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text if cell_text else "")
                
                if any(row_data):  # Only add rows with content
                    table_content.append(f"Row {row_idx + 1}: " + " | ".join(row_data))
            
            if len(table_content) > 1:  # If table has content beyond the title
                table_texts.append("\n".join(table_content))
    
    except Exception as e:
        print(f"Warning: Error extracting tables: {str(e)}")
    
    return table_texts

def extract_document_properties(file_path):
    """
    Extract document properties and metadata (optional enhancement).
    """
    try:
        doc = DocxDocument(file_path)
        properties = doc.core_properties
        
        metadata = {}
        if properties.title:
            metadata['title'] = properties.title
        if properties.author:
            metadata['author'] = properties.author
        if properties.subject:
            metadata['subject'] = properties.subject
        if properties.created:
            metadata['created'] = str(properties.created)
        if properties.modified:
            metadata['modified'] = str(properties.modified)
        
        return metadata
    except:
        return {}